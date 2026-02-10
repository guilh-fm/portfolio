import os
import threading
import win32com.client as win32
from openpyxl import load_workbook
from docx import Document
from docx.shared import RGBColor, Pt
import customtkinter as ctk
from tkinter import filedialog, messagebox

class Pfica:
    def __init__(self, master):
        self.master = master
        
        master.title("PFICA - Preenche Formulários de Interconexão CLARO/ALGAR")
        master.geometry("650x900")
        master.resizable(False, False)
        
        ctk.set_appearance_mode("dark")
        ctk.set_default_color_theme("dark-blue")
        
        self.main_frame = ctk.CTkFrame(master, fg_color="#181818", corner_radius=0)
        self.main_frame.pack(fill="both", expand=True)

        self.caminho_xlsx = ctk.StringVar()
        self.caminho_modelo = ctk.StringVar()
        self.caminho_pasta_saida = ctk.StringVar()
        self.modelo_selecionado = ctk.StringVar(value="CLARO")
        self.prefixo_saida = ctk.StringVar(value="ITX-OPERADORA")

        self.frame_modelo = ctk.CTkFrame(self.main_frame, corner_radius=10, fg_color="white", border_width=2, border_color="#383838")
        self.frame_modelo.pack(pady=10, padx=20, fill="x")
        ctk.CTkLabel(self.frame_modelo, text="1. Selecione a Operadora do Modelo", font=('Arial', 12, 'bold'), text_color="#181818").pack(pady=5, padx=10, anchor="w")

        self.radio_claro = ctk.CTkRadioButton(self.frame_modelo, text="CLARO", variable=self.modelo_selecionado, value="CLARO", fg_color="#383838", hover_color="#7C8B8F", text_color="#181818")
        self.radio_claro.pack(side="left", padx=10, pady=5)
        self.radio_algar = ctk.CTkRadioButton(self.frame_modelo, text="ALGAR", variable=self.modelo_selecionado, value="ALGAR", fg_color="#383838", hover_color="#7C8B8F", text_color="#181818")
        self.radio_algar.pack(side="left", padx=10, pady=5)

        self.frame_selecao_arquivos = ctk.CTkFrame(self.main_frame, corner_radius=10, fg_color="white", border_width=2, border_color="#383838")
        self.frame_selecao_arquivos.pack(pady=10, padx=20, fill="x")
        ctk.CTkLabel(self.frame_selecao_arquivos, text="2. Selecione os Arquivos", font=('Arial', 12, 'bold'), text_color="#181818").pack(pady=5, padx=10, anchor="w")

        ctk.CTkLabel(self.frame_selecao_arquivos, text="Base de Dados - Excel:", text_color="#181818").pack(pady=5, padx=10, anchor="w")
        frame_input_excel = ctk.CTkFrame(self.frame_selecao_arquivos, fg_color="transparent")
        frame_input_excel.pack(fill="x", padx=10, pady=(5,20))
        
        self.entrada_excel = ctk.CTkEntry(frame_input_excel, textvariable=self.caminho_xlsx, width=450, fg_color="#383838", text_color="white")
        self.entrada_excel.pack(side="left", fill="x", expand=True, padx=(0, 5))
        self.btn_excel = ctk.CTkButton(frame_input_excel, text="Procurar...", command=self.procurar_excel, fg_color="#181818", hover_color="#3B3E41", text_color="white")
        self.btn_excel.pack(side="right")

        ctk.CTkLabel(self.frame_selecao_arquivos, text="Formulário - Word:", text_color="#181818").pack(pady=5, padx=10, anchor="w")
        frame_input_modelo = ctk.CTkFrame(self.frame_selecao_arquivos, fg_color="transparent")
        frame_input_modelo.pack(fill="x", padx=10, pady=(5,20))
        
        self.entrada_modelo = ctk.CTkEntry(frame_input_modelo, textvariable=self.caminho_modelo, width=450, fg_color="#383838", text_color="white")
        self.entrada_modelo.pack(side="left", fill="x", expand=True, padx=(0, 5))
        self.btn_modelo = ctk.CTkButton(frame_input_modelo, text="Procurar...", command=self.procurar_modelo, fg_color="#181818", hover_color="#3B3E41", text_color="white")
        self.btn_modelo.pack(side="right")

        self.frame_config_saida = ctk.CTkFrame(self.main_frame, corner_radius=10, fg_color="white", border_width=2, border_color="#383838")
        self.frame_config_saida.pack(pady=10, padx=20, fill="x")
        ctk.CTkLabel(self.frame_config_saida, text="3. Configurações de Saída", font=('Arial', 12, 'bold'), text_color="#181818").pack(pady=5, padx=10, anchor="w")

        ctk.CTkLabel(self.frame_config_saida, text="Prefixo do Nome do Arquivo de Saída:", text_color="#181818").pack(pady=5, padx=10, anchor="w")
        self.entrada_prefixo = ctk.CTkEntry(self.frame_config_saida, textvariable=self.prefixo_saida, fg_color="#383838", text_color="white")
        self.entrada_prefixo.pack(pady=5, padx=10, fill="x", expand=True)
        
        ctk.CTkLabel(self.frame_config_saida, text="Diretório para Salvar Documentos:", text_color="#181818").pack(pady=5, padx=10, anchor="w")
        frame_input_dir = ctk.CTkFrame(self.frame_config_saida, fg_color="transparent")
        frame_input_dir.pack(fill="x", padx=10, pady=(5,20))
        
        self.entrada_dir_saida = ctk.CTkEntry(frame_input_dir, textvariable=self.caminho_pasta_saida, width=450, fg_color="#383838", text_color="white")
        self.entrada_dir_saida.pack(side="left", fill="x", expand=True, padx=(0, 5))
        self.btn_dir_saida = ctk.CTkButton(frame_input_dir, text="Procurar Pasta...", command=self.procurar_pasta_saida, fg_color="#181818", hover_color="#3B3E41", text_color="white")
        self.btn_dir_saida.pack(side="right")

        self.frame_acao = ctk.CTkFrame(self.main_frame, fg_color="transparent")
        self.frame_acao.pack(pady=10, padx=20)
        
        self.btn_gerar = ctk.CTkButton(self.frame_acao, text="Gerar Documentos", command=self.iniciar_thread_geracao, fg_color="#383838", hover_color="#7C8B8F", text_color="white", border_width=2, border_color="white")
        self.btn_gerar.pack(pady=10, fill="x", expand=True)

        self.btn_converter = ctk.CTkButton(self.frame_acao, text="Converter para PDF", command=self.iniciar_thread_conversao, state="disabled", fg_color="#383838", hover_color="#7C8B8F", text_color="white", border_width=2, border_color="white")
        self.btn_converter.pack(pady=5, fill="x", expand=True)
        
        self.frame_status = ctk.CTkFrame(self.main_frame, corner_radius=10, fg_color="white", border_width=2, border_color="#383838")
        self.frame_status.pack(pady=10, padx=20, fill="both", expand=True)
        ctk.CTkLabel(self.frame_status, text="Status do Processo", font=('Arial', 12, 'bold'), text_color="#181818").pack(pady=5, padx=10, anchor="w")
        
        self.barra_progresso = ctk.CTkProgressBar(self.frame_status, orientation="horizontal", mode="determinate", progress_color="#383838", corner_radius=10)
        self.barra_progresso.pack(pady=5, padx=10, fill="x")
        self.barra_progresso.set(0)

        self.texto_status = ctk.CTkTextbox(self.frame_status, height=150, wrap="word", font=('Arial', 11), corner_radius=10, fg_color="#383838", text_color="#181818")
        self.texto_status.pack(pady=10, padx=10, fill="both", expand=True)
        
        self.texto_status.tag_config("sucesso", foreground="green")
        self.texto_status.tag_config("erro", foreground="red")
        self.texto_status.tag_config("info", foreground="white")
        self.texto_status.tag_config("padrao", foreground="#181818")

        self.atualizar_log("Aguardando seleção de arquivos...", "info")

    def atualizar_log(self, mensagem, tag="padrao"):
        self.texto_status.configure(state="normal")
        self.texto_status.insert("end", mensagem + "\n", tag)
        self.texto_status.see("end")
        self.texto_status.configure(state="disabled")
        self.master.update_idletasks()

    def procurar_excel(self):
        arquivo = filedialog.askopenfilename(
            title="Selecione o arquivo Excel (XLSX) com os dados",
            filetypes=(("Arquivos Excel", "*.xlsx"), ("Todos os arquivos", "*.*"))
        )
        if arquivo:
            self.caminho_xlsx.set(arquivo)
            self.atualizar_log(f"Arquivo XLSX selecionado: {os.path.basename(arquivo)}")

    def procurar_modelo(self):
        arquivo = filedialog.askopenfilename(
            title="Selecione o template Word (DOCX) da operadora",
            filetypes=(("Arquivos Word", "*.docx"), ("Todos os arquivos", "*.*"))
        )
        if arquivo:
            self.caminho_modelo.set(arquivo)
            self.atualizar_log(f"Template DOCX selecionado: {os.path.basename(arquivo)}")

    def procurar_pasta_saida(self):
        diretorio = filedialog.askdirectory(
            title="Selecione o diretório para salvar os documentos gerados"
        )
        if diretorio:
            self.caminho_pasta_saida.set(diretorio)
            self.atualizar_log(f"Diretório de saída selecionado: {os.path.basename(diretorio)}")

    def iniciar_thread_geracao(self):
        thread = threading.Thread(target=self.executar_geracao)
        thread.daemon = True
        thread.start()

    def executar_geracao(self):
        arquivo_excel = self.caminho_xlsx.get()
        arquivo_modelo = self.caminho_modelo.get()
        pasta_saida = os.path.normpath(self.caminho_pasta_saida.get())
        prefixo_saida = self.prefixo_saida.get().strip()
        operadora = self.modelo_selecionado.get()

        if not arquivo_excel or not arquivo_modelo or not pasta_saida:
            self.master.after(0, lambda: messagebox.showerror("Erro de Seleção", "Por favor, selecione todos os arquivos e diretórios."))
            return

        self.master.after(0, self.atualizar_log, "Iniciando geração de documentos...", "info")
        self.master.after(0, lambda: self.btn_gerar.configure(state="disabled"))
        self.master.after(0, lambda: self.btn_converter.configure(state="disabled"))

        try:
            dir_docx = os.path.join(pasta_saida, "DOCX")
            dir_pdf = os.path.join(pasta_saida, "PDF")

            if not os.path.exists(dir_docx):
                os.makedirs(dir_docx)
            if not os.path.exists(dir_pdf):
                os.makedirs(dir_pdf)

            self.preencher_documentos_docx(arquivo_excel, arquivo_modelo, dir_docx, prefixo_saida, operadora)
            
            self.master.after(0, self.atualizar_log, "\nGerados todos os arquivos DOCX com sucesso.", "sucesso")
            self.master.after(0, self.atualizar_log, "Pode agora converter para PDF.", "info")
            self.master.after(0, lambda: self.btn_converter.configure(state="normal"))

            self.master.after(0, lambda: messagebox.showinfo("Sucesso", f"Documentos DOCX gerados com sucesso em:\n{dir_docx}"))

        except Exception as e:
            self.master.after(0, self.atualizar_log, f"\nOcorreu um erro: {e}", "erro")
            self.master.after(0, lambda: messagebox.showerror("Erro", f"Ocorreu um erro durante a geração:\n{e}"))
        finally:
            self.master.after(0, lambda: self.btn_gerar.configure(state="normal"))

    def iniciar_thread_conversao(self):
        thread = threading.Thread(target=self.executar_conversao)
        thread.daemon = True
        thread.start()

    def executar_conversao(self):
        pasta_saida = self.caminho_pasta_saida.get()
        if not pasta_saida:
            self.master.after(0, lambda: messagebox.showerror("Erro de Seleção", "Por favor, gere os documentos primeiro."))
            return

        dir_docx = os.path.abspath(os.path.join(pasta_saida, "DOCX"))
        dir_pdf = os.path.abspath(os.path.join(pasta_saida, "PDF"))

        self.master.after(0, self.atualizar_log, "\nIniciando conversão para PDF...", "info")
        self.master.after(0, lambda: self.btn_converter.configure(state="disabled"))
        self.master.after(0, lambda: self.btn_gerar.configure(state="disabled"))

        try:
            self.converter_para_pdf(dir_docx, dir_pdf)
            self.master.after(0, self.atualizar_log, "\nConversão para PDF concluída com sucesso!", "sucesso")
            self.master.after(0, lambda: messagebox.showinfo("Sucesso", f"Todos os arquivos .docx foram convertidos para PDF em:\n{dir_pdf}"))
        except Exception as e:
            self.master.after(0, self.atualizar_log, f"\nOcorreu um erro durante a conversão: {e}", "erro")
            self.master.after(0, lambda: messagebox.showerror("Erro", f"Ocorreu um erro durante a conversão:\n{e}"))
        finally:
            self.master.after(0, lambda: self.btn_converter.configure(state="normal"))
            self.master.after(0, lambda: self.btn_gerar.configure(state="normal"))

    def converter_para_pdf(self, diretorio_origem, diretorio_destino):
        word = None
        try:
            word = win32.Dispatch("Word.Application")
            word.Visible = False

            arquivos = os.listdir(diretorio_origem)
            arquivos_docx = [f for f in arquivos if f.endswith(".docx")]
            total_arquivos = len(arquivos_docx)

            self.master.after(0, lambda: self.barra_progresso.set(0))

            if not arquivos_docx:
                self.master.after(0, self.atualizar_log, "   Nenhum arquivo .docx encontrado para conversão.", "erro")
                return

            for i, nome_arquivo_docx in enumerate(arquivos_docx):
                try:
                    caminho_docx = os.path.join(diretorio_origem, nome_arquivo_docx)
                    nome_pdf = os.path.splitext(nome_arquivo_docx)[0] + ".pdf"
                    caminho_pdf = os.path.join(diretorio_destino, nome_pdf)
                    
                    self.master.after(0, self.atualizar_log, f"   Convertendo {nome_arquivo_docx} para PDF...", "info")
                    
                    doc = word.Documents.Open(caminho_docx)
                    doc.SaveAs(caminho_pdf, FileFormat=17)
                    doc.Close()
                    
                    self.master.after(0, self.atualizar_log, f"   Conversão de {nome_arquivo_docx} concluída!", "padrao")
                    
                    progresso = (i + 1) / total_arquivos
                    self.master.after(0, lambda: self.barra_progresso.set(progresso))

                except Exception as e:
                    self.master.after(0, self.atualizar_log, f"Erro ao converter {nome_arquivo_docx}: {e}", "erro")
                    continue
        except Exception as e:
            self.master.after(0, self.atualizar_log, f"Erro na comunicação com o Word: {e}", "erro")
        finally:
            if word:
                word.Quit()
                word = None
            self.master.after(0, self.atualizar_log, f"   Processo de conversão finalizado.", "sucesso")
    
    def preencher_documentos_docx(self, caminho_planilha_xlsx, caminho_template_docx, diretorio_saida, prefixo_base, operadora_selecionada):
        try:
            pasta_trabalho = load_workbook(filename=caminho_planilha_xlsx)
            aba = pasta_trabalho.active
            total_linhas = aba.max_row - 1
            self.master.after(0, lambda: self.barra_progresso.set(0))

        except Exception as e:
            raise Exception(f"Erro ao carregar a planilha XLSX: {e}")

        mapeamento_colunas = {
            "UF": 0,
            "CN": 1,
            "AREA_LOCAL": 2,
            "SIGLA": 3,
            "MUNICIPIO": 4,
            "ENDERECO": 5,
            "CEP": 6,
            "LATITUDE": 7,
            "LONGITUDE": 8,
            "EOT": 9,
            "PREFIXO": 10,
            "INICIAL": 11,
            "FINAL": 12
        }
        
        nome_base_template_original = os.path.basename(caminho_template_docx)
        nome_template_sem_extensao = os.path.splitext(nome_base_template_original)[0]

        for indice_linha, dados_linha_tupla in enumerate(aba.iter_rows(min_row=2, values_only=True)):
            self.master.after(0, self.atualizar_log, f"  Processando linha {indice_linha + 2}...", "info")
            
            try:
                doc = Document(caminho_template_docx)
            except FileNotFoundError:
                raise FileNotFoundError(f"Template DOCX não encontrado em '{caminho_template_docx}'")
            except Exception as e:
                raise Exception(f"Erro ao carregar o template DOCX: {e}")

            dados_linha = {}
            for marcador_codigo, col_index in mapeamento_colunas.items():
                valor = dados_linha_tupla[col_index] if col_index < len(dados_linha_tupla) else ""
                dados_linha[marcador_codigo] = str(valor) if valor is not None else ""

            def substituir_e_formatar_paragrafo(paragrafo):
                texto_completo_atual = "".join([run.text for run in paragrafo.runs])
                texto_modificado = texto_completo_atual
                
                formato_base_run = None
                if paragrafo.runs:
                    formato_base_run = paragrafo.runs[0].font
                
                houve_substituicao = False
                for marcador_codigo, valor_planilha in dados_linha.items():
                    marcador_template = f"{{{{{marcador_codigo}}}}}"

                    if marcador_template in texto_modificado:
                        texto_modificado = texto_modificado.replace(marcador_template, valor_planilha)
                        houve_substituicao = True

                if houve_substituicao:

                    for run_para_remover in list(paragrafo.runs):
                        paragrafo._element.remove(run_para_remover._element)

                    novo_run = paragrafo.add_run(texto_modificado)
                    
                    if formato_base_run:
                        novo_run.font.name = formato_base_run.name
                        novo_run.font.size = formato_base_run.size
                        novo_run.font.bold = formato_base_run.bold
                        novo_run.font.italic = formato_base_run.italic
                        novo_run.font.underline = formato_base_run.underline
                        novo_run.font.color.rgb = RGBColor(0, 0, 0)
                    else:
                        novo_run.font.name = 'Arial Narrow'
                        novo_run.font.size = Pt(10)
                        novo_run.font.color.rgb = RGBColor(0, 0, 0)

            for p in doc.paragraphs:
                substituir_e_formatar_paragrafo(p)

            for tabela in doc.tables:
                for linha_tabela in tabela.rows:
                    for celula in linha_tabela.cells:
                        for p_na_celula in celula.paragraphs:
                            substituir_e_formatar_paragrafo(p_na_celula)

            nome_base_arquivo_saida = ""
            if prefixo_base:
                nome_base_arquivo_saida = f"{prefixo_base}-{operadora_selecionada}"
            else:
                temp_nome = nome_template_sem_extensao
                if " - Base" in temp_nome:
                    temp_nome = temp_nome.replace(" - Base", "").strip()
                elif " - BASE" in temp_nome:
                    temp_nome = temp_nome.replace(" - BASE", "").strip()
                
                if operadora_selecionada.upper() not in temp_nome.upper():
                    nome_base_arquivo_saida = f"{temp_nome}-{operadora_selecionada}"
                else:
                    nome_base_arquivo_saida = temp_nome

            nome_arquivo_saida = (
                f"{nome_base_arquivo_saida} - {dados_linha.get('SIGLA', 'Sem_Sigla')}.docx"
            )
            caminho_arquivo_saida = os.path.join(diretorio_saida, nome_arquivo_saida)

            try:
                doc.save(caminho_arquivo_saida)
                self.master.after(0, self.atualizar_log, f"  Gerado: {os.path.basename(caminho_arquivo_saida)}")
                
                progresso = (indice_linha + 1) / total_linhas
                self.master.after(0, lambda: self.barra_progresso.set(progresso))
            except Exception as e:
                raise Exception(f"Erro ao salvar o documento '{nome_arquivo_saida}': {e}")

if __name__ == "__main__":
    raiz = ctk.CTk()
    app = Pfica(raiz)
    raiz.mainloop()