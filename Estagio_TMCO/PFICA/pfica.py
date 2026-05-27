import os
import threading
from tkinter import filedialog, messagebox

import customtkinter as ctk
import win32com.client as win32
from docx import Document
from docx.shared import Pt, RGBColor
from openpyxl import load_workbook


MAPEAMENTO_COLUNAS = {
    "UF": 0,
    "CN": 1,
    "AREA_LOCAL": 2,
    "SIGLA": 3,
    "MUNICIPIO": 4,
    "ENDERECO": 5,
    "CEP": 6,
    "LATITUDE": 7,
    "LONGITUDE": 8,
    "EOT": 9,
    "PREFIXO": 10,
    "INICIAL": 11,
    "FINAL": 12,
}


# Funções auxiliares

def criar_pastas_saida(pasta_base):
    pasta_docx = os.path.join(pasta_base, "DOCX")
    pasta_pdf = os.path.join(pasta_base, "PDF")

    if not os.path.exists(pasta_docx):
        os.makedirs(pasta_docx)

    if not os.path.exists(pasta_pdf):
        os.makedirs(pasta_pdf)

    return pasta_docx, pasta_pdf


def validar_campos_obrigatorios(caminho_excel, caminho_modelo, pasta_saida):
    if caminho_excel == "" or caminho_modelo == "" or pasta_saida == "":
        return False

    return True


def carregar_planilha(caminho_planilha):
    try:
        pasta_trabalho = load_workbook(filename=caminho_planilha, data_only=True)
        aba = pasta_trabalho.active
        return aba
    except Exception as erro:
        raise Exception(f"Erro ao carregar a planilha de dados: {erro}")


def carregar_documento_modelo(caminho_modelo):
    try:
        documento = Document(caminho_modelo)
        return documento
    except FileNotFoundError:
        raise FileNotFoundError(f"Modelo DOCX não encontrado em: {caminho_modelo}")
    except Exception as erro:
        raise Exception(f"Erro ao carregar o modelo DOCX: {erro}")


def montar_dados_linha(dados_linha):
    dados = {}

    for marcador, indice_coluna in MAPEAMENTO_COLUNAS.items():
        if indice_coluna < len(dados_linha) and dados_linha[indice_coluna] is not None:
            dados[marcador] = str(dados_linha[indice_coluna])
        else:
            dados[marcador] = ""

    return dados


def copiar_formatacao_run(novo_run, formato_base):
    if formato_base is None:
        novo_run.font.name = "Arial Narrow"
        novo_run.font.size = Pt(10)
        novo_run.font.color.rgb = RGBColor(0, 0, 0)
        return

    novo_run.font.name = formato_base.name
    novo_run.font.size = formato_base.size
    novo_run.font.bold = formato_base.bold
    novo_run.font.italic = formato_base.italic
    novo_run.font.underline = formato_base.underline
    novo_run.font.color.rgb = RGBColor(0, 0, 0)


def substituir_marcadores_paragrafo(paragrafo, dados_linha):
    texto_original = "".join(run.text for run in paragrafo.runs)
    texto_modificado = texto_original

    for marcador, valor in dados_linha.items():
        tag = "{{" + marcador + "}}"
        texto_modificado = texto_modificado.replace(tag, valor)

    if texto_modificado == texto_original:
        return

    if paragrafo.runs:
        formato_base = paragrafo.runs[0].font
    else:
        formato_base = None

    for run in list(paragrafo.runs):
        paragrafo._element.remove(run._element)

    novo_run = paragrafo.add_run(texto_modificado)
    copiar_formatacao_run(novo_run, formato_base)


def substituir_marcadores_documento(documento, dados_linha):
    for paragrafo in documento.paragraphs:
        substituir_marcadores_paragrafo(paragrafo, dados_linha)

    for tabela in documento.tables:
        for linha_tabela in tabela.rows:
            for celula in linha_tabela.cells:
                for paragrafo in celula.paragraphs:
                    substituir_marcadores_paragrafo(paragrafo, dados_linha)


def limpar_nome_template(nome_template):
    nome_limpo = nome_template

    if " - Base" in nome_limpo:
        nome_limpo = nome_limpo.replace(" - Base", "").strip()
    elif " - BASE" in nome_limpo:
        nome_limpo = nome_limpo.replace(" - BASE", "").strip()

    return nome_limpo


def montar_nome_base_saida(caminho_modelo, prefixo_base, operadora):
    if prefixo_base != "":
        return f"{prefixo_base}-{operadora}"

    nome_template = os.path.basename(caminho_modelo)
    nome_sem_extensao = os.path.splitext(nome_template)[0]
    nome_limpo = limpar_nome_template(nome_sem_extensao)

    if operadora.upper() not in nome_limpo.upper():
        return f"{nome_limpo}-{operadora}"

    return nome_limpo


def montar_nome_arquivo_docx(nome_base_saida, dados_linha):
    sigla = dados_linha.get("SIGLA", "")

    if sigla == "":
        sigla = "Sem_Sigla"

    nome_arquivo = f"{nome_base_saida} - {sigla}.docx"
    return nome_arquivo


def salvar_documento(documento, diretorio_saida, nome_arquivo):
    caminho_saida = os.path.join(diretorio_saida, nome_arquivo)
    documento.save(caminho_saida)
    return caminho_saida


def listar_arquivos_docx(diretorio_origem):
    arquivos = os.listdir(diretorio_origem)
    arquivos_docx = []

    for nome_arquivo in arquivos:
        if nome_arquivo.lower().endswith(".docx"):
            arquivos_docx.append(nome_arquivo)

    return arquivos_docx


# Interface gráfica

class AplicacaoPfica:
    def __init__(self, janela):
        self.janela = janela

        self.janela.title("PFICA - Preenche Formulários de Interconexão CLARO/ALGAR")
        self.janela.geometry("650x900")
        self.janela.resizable(False, False)

        ctk.set_appearance_mode("dark")
        ctk.set_default_color_theme("dark-blue")

        self.caminho_excel = ctk.StringVar()
        self.caminho_modelo = ctk.StringVar()
        self.caminho_pasta_saida = ctk.StringVar()
        self.operadora_selecionada = ctk.StringVar(value="CLARO")
        self.prefixo_saida = ctk.StringVar(value="ITX-OPERADORA")

        self.criar_componentes()
        self.atualizar_log("Aguardando seleção de arquivos...", "info")

    def criar_componentes(self):
        self.frame_principal = ctk.CTkFrame(self.janela, fg_color="#181818", corner_radius=0)
        self.frame_principal.pack(fill="both", expand=True)

        self.criar_bloco_operadora()
        self.criar_bloco_arquivos()
        self.criar_bloco_saida()
        self.criar_bloco_acoes()
        self.criar_bloco_status()

    def criar_bloco_operadora(self):
        frame = ctk.CTkFrame(self.frame_principal, corner_radius=10, fg_color="white", border_width=2, border_color="#383838")
        frame.pack(pady=10, padx=20, fill="x")

        ctk.CTkLabel(frame, text="1. Selecione a Operadora do Modelo", font=("Arial", 12, "bold"), text_color="#181818").pack(pady=5, padx=10, anchor="w")

        radio_claro = ctk.CTkRadioButton(frame, text="CLARO", variable=self.operadora_selecionada, value="CLARO", fg_color="#383838", hover_color="#7C8B8F", text_color="#181818")
        radio_algar = ctk.CTkRadioButton(frame, text="ALGAR", variable=self.operadora_selecionada, value="ALGAR", fg_color="#383838", hover_color="#7C8B8F", text_color="#181818")

        radio_claro.pack(side="left", padx=10, pady=5)
        radio_algar.pack(side="left", padx=10, pady=5)

    def criar_bloco_arquivos(self):
        frame = ctk.CTkFrame(self.frame_principal, corner_radius=10, fg_color="white", border_width=2, border_color="#383838")
        frame.pack(pady=10, padx=20, fill="x")

        ctk.CTkLabel(frame, text="2. Selecione os Arquivos", font=("Arial", 12, "bold"), text_color="#181818").pack(pady=5, padx=10, anchor="w")
        self.criar_linha_arquivo(frame, "Base de Dados - Excel:", self.caminho_excel, self.procurar_excel)
        self.criar_linha_arquivo(frame, "Formulário - Word:", self.caminho_modelo, self.procurar_modelo)

    def criar_bloco_saida(self):
        frame = ctk.CTkFrame(self.frame_principal, corner_radius=10, fg_color="white", border_width=2, border_color="#383838")
        frame.pack(pady=10, padx=20, fill="x")

        ctk.CTkLabel(frame, text="3. Configurações de Saída", font=("Arial", 12, "bold"), text_color="#181818").pack(pady=5, padx=10, anchor="w")
        ctk.CTkLabel(frame, text="Prefixo do Nome do Arquivo de Saída:", text_color="#181818").pack(pady=5, padx=10, anchor="w")

        entrada_prefixo = ctk.CTkEntry(frame, textvariable=self.prefixo_saida, fg_color="#383838", text_color="white")
        entrada_prefixo.pack(pady=5, padx=10, fill="x")

        self.criar_linha_arquivo(frame, "Diretório para Salvar Documentos:", self.caminho_pasta_saida, self.procurar_pasta_saida, texto_botao="Procurar Pasta...")

    def criar_bloco_acoes(self):
        frame = ctk.CTkFrame(self.frame_principal, fg_color="transparent")
        frame.pack(pady=10, padx=20)

        self.botao_gerar = ctk.CTkButton(
            frame,
            text="Gerar Documentos",
            command=self.iniciar_thread_geracao,
            fg_color="#383838",
            hover_color="#7C8B8F",
            text_color="white",
            border_width=2,
            border_color="white",
        )
        self.botao_gerar.pack(pady=10, fill="x")

        self.botao_converter = ctk.CTkButton(
            frame,
            text="Converter para PDF",
            command=self.iniciar_thread_conversao,
            state="disabled",
            fg_color="#383838",
            hover_color="#7C8B8F",
            text_color="white",
            border_width=2,
            border_color="white",
        )
        self.botao_converter.pack(pady=5, fill="x")

    def criar_bloco_status(self):
        frame = ctk.CTkFrame(self.frame_principal, corner_radius=10, fg_color="white", border_width=2, border_color="#383838")
        frame.pack(pady=10, padx=20, fill="both", expand=True)

        ctk.CTkLabel(frame, text="Status do Processo", font=("Arial", 12, "bold"), text_color="#181818").pack(pady=5, padx=10, anchor="w")

        self.barra_progresso = ctk.CTkProgressBar(frame, orientation="horizontal", mode="determinate", progress_color="#383838", corner_radius=10)
        self.barra_progresso.pack(pady=5, padx=10, fill="x")
        self.barra_progresso.set(0)

        self.texto_status = ctk.CTkTextbox(frame, height=150, wrap="word", font=("Arial", 11), corner_radius=10, fg_color="#383838", text_color="white")
        self.texto_status.pack(pady=10, padx=10, fill="both", expand=True)

        self.texto_status.tag_config("sucesso", foreground="green")
        self.texto_status.tag_config("erro", foreground="red")
        self.texto_status.tag_config("info", foreground="white")
        self.texto_status.tag_config("padrao", foreground="white")

    def criar_linha_arquivo(self, frame_pai, texto_label, variavel, comando, texto_botao="Procurar..."):
        ctk.CTkLabel(frame_pai, text=texto_label, text_color="#181818").pack(pady=5, padx=10, anchor="w")

        frame_input = ctk.CTkFrame(frame_pai, fg_color="transparent")
        frame_input.pack(fill="x", padx=10, pady=(5, 20))

        entrada = ctk.CTkEntry(frame_input, textvariable=variavel, width=450, fg_color="#383838", text_color="white")
        entrada.pack(side="left", fill="x", expand=True, padx=(0, 5))

        botao = ctk.CTkButton(frame_input, text=texto_botao, command=comando, fg_color="#181818", hover_color="#3B3E41", text_color="white")
        botao.pack(side="right")

    def atualizar_log(self, mensagem, tag="padrao"):
        self.texto_status.configure(state="normal")
        self.texto_status.insert("end", mensagem + "\n", tag)
        self.texto_status.see("end")
        self.texto_status.configure(state="disabled")
        self.janela.update_idletasks()

    def procurar_excel(self):
        arquivo = filedialog.askopenfilename(
            title="Selecione o arquivo Excel com os dados",
            filetypes=(("Arquivos Excel", "*.xlsx"), ("Todos os arquivos", "*.*")),
        )

        if arquivo:
            self.caminho_excel.set(arquivo)
            self.atualizar_log(f"Arquivo selecionado: {os.path.basename(arquivo)}")

    def procurar_modelo(self):
        arquivo = filedialog.askopenfilename(
            title="Selecione o modelo Word da operadora",
            filetypes=(("Arquivos Word", "*.docx"), ("Todos os arquivos", "*.*")),
        )

        if arquivo:
            self.caminho_modelo.set(arquivo)
            self.atualizar_log(f"Modelo selecionado: {os.path.basename(arquivo)}")

    def procurar_pasta_saida(self):
        diretorio = filedialog.askdirectory(title="Selecione o diretório para salvar os documentos")

        if diretorio:
            self.caminho_pasta_saida.set(diretorio)
            self.atualizar_log(f"Diretório de saída selecionado: {os.path.basename(diretorio)}")

    def iniciar_thread_geracao(self):
        thread = threading.Thread(target=self.executar_geracao)
        thread.daemon = True
        thread.start()

    def executar_geracao(self):
        caminho_excel = self.caminho_excel.get()
        caminho_modelo = self.caminho_modelo.get()
        pasta_saida = os.path.normpath(self.caminho_pasta_saida.get())
        prefixo_saida = self.prefixo_saida.get().strip()
        operadora = self.operadora_selecionada.get()

        if not validar_campos_obrigatorios(caminho_excel, caminho_modelo, pasta_saida):
            self.janela.after(0, lambda: messagebox.showerror("Erro de Seleção", "Por favor, selecione todos os arquivos e diretórios."))
            return

        self.janela.after(0, self.atualizar_log, "Iniciando geração de documentos...", "info")
        self.janela.after(0, lambda: self.botao_gerar.configure(state="disabled"))
        self.janela.after(0, lambda: self.botao_converter.configure(state="disabled"))

        try:
            pasta_docx, _ = criar_pastas_saida(pasta_saida)
            self.preencher_documentos(caminho_excel, caminho_modelo, pasta_docx, prefixo_saida, operadora)

            self.janela.after(0, self.atualizar_log, "\nTodos os arquivos DOCX foram gerados com sucesso.", "sucesso")
            self.janela.after(0, self.atualizar_log, "Agora você pode converter os documentos para PDF.", "info")
            self.janela.after(0, lambda: self.botao_converter.configure(state="normal"))
            self.janela.after(0, lambda: messagebox.showinfo("Sucesso", f"Documentos DOCX gerados em:\n{pasta_docx}"))
        except Exception as erro:
            mensagem_erro = str(erro)
            self.janela.after(0, self.atualizar_log, f"\nOcorreu um erro: {mensagem_erro}", "erro")
            self.janela.after(0, lambda: messagebox.showerror("Erro", f"Ocorreu um erro durante a geração:\n{mensagem_erro}"))
        finally:
            self.janela.after(0, lambda: self.botao_gerar.configure(state="normal"))

    def preencher_documentos(self, caminho_planilha, caminho_modelo, diretorio_saida, prefixo_saida, operadora):
        aba_dados = carregar_planilha(caminho_planilha)
        linhas_dados = list(aba_dados.iter_rows(min_row=2, values_only=True))
        linhas_validas = [linha for linha in linhas_dados if any(linha)]

        if len(linhas_validas) == 0:
            raise Exception("A planilha de dados está vazia ou não possui dados após o cabeçalho.")

        nome_base_saida = montar_nome_base_saida(caminho_modelo, prefixo_saida, operadora)
        self.janela.after(0, lambda: self.barra_progresso.set(0))

        for indice, dados_linha_tupla in enumerate(linhas_validas):
            numero_linha = indice + 2
            self.janela.after(0, self.atualizar_log, f"  Processando linha {numero_linha}...", "info")

            dados_linha = montar_dados_linha(dados_linha_tupla)
            documento = carregar_documento_modelo(caminho_modelo)

            substituir_marcadores_documento(documento, dados_linha)

            nome_arquivo = montar_nome_arquivo_docx(nome_base_saida, dados_linha)
            caminho_saida = salvar_documento(documento, diretorio_saida, nome_arquivo)

            self.janela.after(0, self.atualizar_log, f"  Gerado: {os.path.basename(caminho_saida)}", "padrao")

            progresso = (indice + 1) / len(linhas_validas)
            self.janela.after(0, lambda valor=progresso: self.barra_progresso.set(valor))

    def iniciar_thread_conversao(self):
        thread = threading.Thread(target=self.executar_conversao)
        thread.daemon = True
        thread.start()

    def executar_conversao(self):
        pasta_saida = self.caminho_pasta_saida.get()

        if pasta_saida == "":
            self.janela.after(0, lambda: messagebox.showerror("Erro de Seleção", "Por favor, gere os documentos primeiro."))
            return

        pasta_docx = os.path.abspath(os.path.join(pasta_saida, "DOCX"))
        pasta_pdf = os.path.abspath(os.path.join(pasta_saida, "PDF"))

        self.janela.after(0, self.atualizar_log, "\nIniciando conversão para PDF...", "info")
        self.janela.after(0, lambda: self.botao_converter.configure(state="disabled"))
        self.janela.after(0, lambda: self.botao_gerar.configure(state="disabled"))

        try:
            self.converter_para_pdf(pasta_docx, pasta_pdf)
            self.janela.after(0, self.atualizar_log, "\nConversão para PDF concluída com sucesso.", "sucesso")
            self.janela.after(0, lambda: messagebox.showinfo("Sucesso", f"Arquivos convertidos para PDF em:\n{pasta_pdf}"))
        except Exception as erro:
            mensagem_erro = str(erro)
            self.janela.after(0, self.atualizar_log, f"\nOcorreu um erro durante a conversão: {mensagem_erro}", "erro")
            self.janela.after(0, lambda: messagebox.showerror("Erro", f"Ocorreu um erro durante a conversão:\n{mensagem_erro}"))
        finally:
            self.janela.after(0, lambda: self.botao_converter.configure(state="normal"))
            self.janela.after(0, lambda: self.botao_gerar.configure(state="normal"))

    def converter_para_pdf(self, diretorio_origem, diretorio_destino):
        word = None

        try:
            word = win32.Dispatch("Word.Application")
            word.Visible = False

            arquivos_docx = listar_arquivos_docx(diretorio_origem)

            if len(arquivos_docx) == 0:
                self.janela.after(0, self.atualizar_log, "  Nenhum arquivo DOCX foi encontrado para conversão.", "erro")
                return

            self.janela.after(0, lambda: self.barra_progresso.set(0))

            for indice, nome_arquivo in enumerate(arquivos_docx):
                caminho_docx = os.path.join(diretorio_origem, nome_arquivo)
                nome_pdf = os.path.splitext(nome_arquivo)[0] + ".pdf"
                caminho_pdf = os.path.join(diretorio_destino, nome_pdf)

                self.janela.after(0, self.atualizar_log, f"  Convertendo {nome_arquivo} para PDF...", "info")

                documento = word.Documents.Open(caminho_docx)
                documento.SaveAs(caminho_pdf, FileFormat=17)
                documento.Close()

                self.janela.after(0, self.atualizar_log, f"  Convertido: {nome_pdf}", "padrao")

                progresso = (indice + 1) / len(arquivos_docx)
                self.janela.after(0, lambda valor=progresso: self.barra_progresso.set(valor))
        except Exception as erro:
            raise Exception(f"Erro na comunicação com o Microsoft Word: {erro}")
        finally:
            if word is not None:
                word.Quit()


if __name__ == "__main__":
    raiz = ctk.CTk()
    app = AplicacaoPfica(raiz)
    raiz.mainloop()
